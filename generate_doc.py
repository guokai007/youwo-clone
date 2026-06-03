#!/usr/bin/env python3
"""生成游我克隆项目 Word 文档"""
import sys
sys.path.insert(0, '/Users/guokai/youwo-clone/.venv/lib/python3.11/site-packages')

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = '/Users/guokai/youwo-clone'
FILL = 'f7fffe'


def set_bg(p):
    """给段落加浅绿背景"""
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), FILL)
    pPr.append(shd)


def add_code(doc, code):
    for line in code.split('\n'):
        p = doc.add_paragraph()
        r = p.add_run(line if line else ' ')
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0x2e, 0x86, 0x4a)
        set_bg(p)


def add_file(doc, path, rel_path):
    h = doc.add_heading(f'  {rel_path}', level=3)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x31, 0x5b, 0x59)
        run.font.size = Pt(10)

    try:
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
    except Exception as e:
        doc.add_paragraph(f'[文件读取失败: {e}]')
        return

    add_code(doc, content)
    doc.add_paragraph()


def main():
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # ===== 封面 =====
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run('游我克隆项目')
    tr.font.size = Pt(28)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(0x31, 0x5b, 0x59)

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run('—— 从技术解析到完整实现（实验版）\n包含全部源代码')
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(0x69, 0x7b, 0x7a)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run(f'生成时间：2026-06-03  |  项目路径：~/youwo-clone/')
    ir.font.size = Pt(10)
    ir.font.color.rgb = RGBColor(0x9a, 0x9a, 0x9a)

    doc.add_page_break()

    # ===== 项目文件清单 =====
    doc.add_heading('项目文件清单', level=1)

    files = [
        ('前端项目配置', 'frontend/package.json'),
        ('Vite 构建配置', 'frontend/vite.config.js'),
        ('HTML 入口', 'frontend/index.html'),
        ('前端入口', 'frontend/src/main.js'),
        ('根组件', 'frontend/src/App.vue'),
        ('路由配置', 'frontend/src/router/index.js'),
        ('聊天状态管理', 'frontend/src/stores/chat.js'),
        ('主样式', 'frontend/src/assets/main.css'),
        ('今天首页', 'frontend/src/views/Today.vue'),
        ('聊聊对话页', 'frontend/src/views/Chat.vue'),
        ('留存记录页', 'frontend/src/views/Records.vue'),
        ('回望回忆页', 'frontend/src/views/Review.vue'),
        ('偏好设置页', 'frontend/src/views/Preferences.vue'),
        ('Python 依赖', 'backend/requirements.txt'),
        ('FastAPI 主入口', 'backend/main.py'),
        ('聊天 API 路由', 'backend/routers/chat.py'),
        ('记忆 API 路由', 'backend/routers/memory.py'),
        ('AI 模型调用', 'backend/services/llm.py'),
        ('环境变量示例', 'backend/.env.example'),
        ('部署脚本', 'deploy/deploy.sh'),
        ('Docker 配置', 'deploy/docker-compose.yml'),
        ('Nginx 配置', 'deploy/nginx.conf'),
    ]

    for desc, path in files:
        full = os.path.join(BASE, path)
        p = doc.add_paragraph()
        p.add_run(f'• {desc}：').bold = True
        p.add_run(path)
        p.style = doc.styles['List Bullet']

    doc.add_page_break()

    # ===== 各文件内容 =====
    sections_files = [
        ('一、前端项目配置', [
            ('frontend/package.json', '前端依赖配置'),
            ('frontend/vite.config.js', 'Vite 构建配置'),
            ('frontend/index.html', 'HTML 入口'),
        ]),
        ('二、前端核心文件', [
            ('frontend/src/main.js', '前端入口文件'),
            ('frontend/src/App.vue', '根组件（侧边栏+路由视图）'),
            ('frontend/src/router/index.js', 'Vue Router 配置'),
            ('frontend/src/stores/chat.js', 'Pinia 聊天状态管理'),
            ('frontend/src/assets/main.css', '完整样式表（复刻游我风格）'),
        ]),
        ('三、前端页面组件', [
            ('frontend/src/views/Today.vue', '今天首页'),
            ('frontend/src/views/Chat.vue', '聊聊对话页（核心）'),
            ('frontend/src/views/Records.vue', '留存记录页'),
            ('frontend/src/views/Review.vue', '回望回忆页'),
            ('frontend/src/views/Preferences.vue', '偏好设置页'),
        ]),
        ('四、后端服务', [
            ('backend/requirements.txt', 'Python 依赖'),
            ('backend/main.py', 'FastAPI 主入口'),
            ('backend/routers/chat.py', '聊天 API 路由'),
            ('backend/routers/memory.py', '记忆 API 路由'),
            ('backend/services/llm.py', 'AI 大模型调用服务'),
            ('backend/.env.example', '环境变量配置示例'),
        ]),
        ('五、部署脚本', [
            ('deploy/deploy.sh', 'VPS 一键部署脚本'),
            ('deploy/docker-compose.yml', 'Docker Compose 编排'),
            ('deploy/nginx.conf', 'Nginx 配置'),
        ]),
    ]

    for section_title, section_files in sections_files:
        doc.add_heading(section_title, level=1)
        for file_path, desc in section_files:
            full_path = os.path.join(BASE, file_path)
            if os.path.exists(full_path):
                p = doc.add_paragraph()
                p.add_run(f'📄 {file_path} —— {desc}').bold = True
                p.runs[0].font.color.rgb = RGBColor(0x31, 0x5b, 0x59)
                add_code(doc, open(full_path, encoding='utf-8').read())
                doc.add_paragraph()
            else:
                doc.add_paragraph(f'[文件不存在] {file_path}')
        doc.add_page_break()

    # ===== 部署说明 =====
    doc.add_heading('六、本地运行与部署说明', level=1)

    doc.add_heading('1. 前端开发模式', level=2)
    add_code(doc, '''cd ~/youwo-clone/frontend
npm install
npm run dev
# 访问 http://localhost:5173''')

    doc.add_heading('2. 后端启动', level=2)
    add_code(doc, '''cd ~/youwo-clone/backend
cp .env.example .env
# 编辑 .env，填入 AI API Key
uvicorn main:app --host 127.0.0.1 --port 8000 --reload''')

    doc.add_heading('3. 前端构建', level=2)
    add_code(doc, '''cd ~/youwo-clone/frontend
npm run build
# 构建产物在 dist/ 目录''')

    doc.add_heading('4. AI API 配置（backend/.env）', level=2)
    add_code(doc, '''# 方式1: MiniMax（推荐国内用户）
LLM_API_KEY=your_api_key_here

# 方式2: 硅基流动（便宜稳定，兼容 OpenAI）
# LLM_API_KEY=your_api_key_here
# LLM_BASE_URL=https://api.siliconflow.cn/v1
# LLM_MODEL=THUDM/glm-4-flash

# 方式3: OpenAI（需要代理）
# LLM_API_KEY=sk-...
# LLM_BASE_URL=https://api.openai.com/v1
# LLM_MODEL=gpt-4o-mini''')

    doc.add_heading('5. VPS 部署（一键脚本）', level=2)
    add_code(doc, '''# 在服务器上安装 SSH 后，设置环境变量并运行
export SERVER_IP=你的服务器IP
export DEPLOY_USER=root
export DOMAIN=你的域名

bash ~/youwo-clone/deploy/deploy.sh''')

    doc.add_heading('6. Docker 部署（可选）', level=2)
    add_code(doc, '''cd ~/youwo-clone/deploy
docker-compose up -d
# 访问 http://服务器IP''')

    doc.add_heading('7. 技术架构', level=2)
    arch = doc.add_paragraph()
    arch.add_run('前端：').bold = True
    arch.add_run('Vue 3 + Vite + Pinia + Vue Router，Hash 路由，localStorage 本地存储\n')
    arch.add_run('后端：').bold = True
    arch.add_run('Python FastAPI + Uvicorn，支持 MiniMax / 硅基流动 / OpenAI\n')
    arch.add_run('样式：').bold = True
    arch.add_run('原生 CSS（变量 + Grid/Flexbox），复刻游我文艺风格\n')
    arch.add_run('字体：').bold = True
    arch.add_run('霞鹜文楷（LXGW WenKai）+ 思源宋体\n')
    arch.add_run('部署：').bold = True
    arch.add_run('Nginx + systemd（VPS）或 Docker Compose\n')
    arch.add_run('服务器：').bold = True
    arch.add_run('可用阿里云/腾讯云/火山引擎等任意 VPS')

    # ===== 保存 =====
    output = '/Users/guokai/游我克隆项目文档.docx'
    doc.save(output)
    print(f'文档已保存: {output}')
    print(f'文件大小: {os.path.getsize(output) / 1024:.1f} KB')


if __name__ == '__main__':
    main()
